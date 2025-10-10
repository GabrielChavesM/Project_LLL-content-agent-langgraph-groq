import streamlit as st
import os
import operator
import random
import re
from typing import TypedDict, Annotated
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from io import BytesIO

# --- LangChain Imports ---
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.messages import HumanMessage, SystemMessage

# --- LangGraph Imports ---
from langgraph.graph import StateGraph, END

# ==========================================================
# I. CONFIGURAÇÃO DE AMBIENTE E LLM
# ==========================================================

load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
LANGCHAIN_API_KEY = os.getenv("LANGCHAIN_API_KEY")

llm = ChatGroq(
    model_name="llama-3.1-8b-instant",
    temperature=0.7,
    groq_api_key=GROQ_API_KEY
)

# ==========================================================
# II. ESTADO DO GRAFO
# ==========================================================

class CVAgentState(TypedDict):
    cv_text: str
    job_description: str
    profile: str
    experiences: str
    education: str
    skills: str
    projects: str
    certifications: str
    plan: Annotated[str, operator.add]
    optimized_cv: Annotated[str, operator.add]
    compatibility: float
    review_feedback: str
    revision_needed: bool
    revision_count: int

# ==========================================================
# III. NÓS DO GRAFO
# ==========================================================

def evaluate_compatibility(state: CVAgentState):
    st.info("📊 Avaliando compatibilidade entre o CV e a vaga...")
    cv_text = state["cv_text"]
    job_desc = state["job_description"]

    prompt = ChatPromptTemplate.from_messages([
        SystemMessage(content=(
            "Você é um avaliador de recrutamento especializado em matching de perfis. "
            "Analise a compatibilidade entre o currículo e a descrição da vaga e dê uma nota percentual (0-100). "
            "Apenas retorne a porcentagem, sem explicações adicionais."
        )),
        HumanMessage(content=f"CURRÍCULO:\n{cv_text}\n\nDESCRIÇÃO DA VAGA:\n{job_desc}")
    ])

    chain = prompt | llm
    result = chain.invoke({}).content.strip()

    try:
        compatibility = float(result.replace("%", "").strip())
    except:
        compatibility = 70.0

    return {"compatibility": compatibility}

def analyze_cv(state: CVAgentState):
    st.info("🧭 Analisando currículo e descrição da vaga...")
    cv_text = state["cv_text"]
    job_desc = state["job_description"]

    full_cv_context = f"""
PERFIL: {state.get('profile', '')}
EXPERIÊNCIAS: {state.get('experiences', '')}
EDUCAÇÃO: {state.get('education', '')}
HABILIDADES: {state.get('skills', '')}
PROJETOS: {state.get('projects', '')}
CERTIFICAÇÕES: {state.get('certifications', '')}
CV ORIGINAL:
{cv_text}
"""

    prompt = ChatPromptTemplate.from_messages([
        SystemMessage(content=(
            "Você é um especialista em RH e recrutamento. "
            "Analise o currículo completo e a vaga e gere um plano detalhado de otimização com foco em aumentar a compatibilidade. "
            "Liste sugestões práticas e palavras-chave relevantes."
        )),
        HumanMessage(content=f"{full_cv_context}\n\nDESCRIÇÃO DA VAGA:\n{job_desc}")
    ])

    chain = prompt | llm
    plan_result = chain.invoke({}).content

    return {"plan": plan_result, "optimized_cv": "", "review_feedback": "", "revision_count": 0}

def optimize_cv(state: CVAgentState):
    st.info("✏️ Reescrevendo currículo adaptado à vaga...")
    plan = state["plan"]
    cv_text = state["cv_text"]
    job_desc = state["job_description"]
    feedback = state.get("review_feedback", "")
    revision_count = state.get("revision_count", 0) + 1

    # --- detectar idioma da vaga (português ou inglês) ---
    language_prompt = ChatPromptTemplate.from_messages([
        SystemMessage(content="Você é um detector de idioma. Responda apenas 'pt' para português ou 'en' para inglês."),
        HumanMessage(content=job_desc)
    ])
    detected_lang = (language_prompt | llm).invoke({}).content.strip().lower()
    if detected_lang not in ["pt", "en"]:
        detected_lang = "pt"

    # Ordem aleatória das seções técnicas
    sections_order = ["Habilidades Técnicas", "Linguagens de Programação", "Conceitos"]
    random.shuffle(sections_order)

    structure_description = f"""
Reescreva o currículo baseado no plano e nas instruções a seguir:
- Idioma: o mesmo da descrição da vaga ({'Português' if detected_lang == 'pt' else 'Inglês'}).
- O currículo deve conter no máximo **500 palavras**.
- Não deve conter erros ortográficos ou gramaticais.
- Estrutura fixa e única:
  1. Profile/Summary
  2. Experiência
  3. Educação
  4. {', '.join(sections_order)}
  5. Projetos
  6. Certificações
- Gere apenas UMA ocorrência de cada seção.
- Os títulos das seções devem estar em NEGRITO, sem markdown.
- Cabeçalho deve conter: LinkedIn, GitHub, site, email, telefone.
- Incluir um espaço reservado para a foto do candidato (esquerda ou direita, aleatoriamente).
- Não ultrapasse 2 páginas (~1000 linhas) e evite redundâncias.
"""

    prompt = ChatPromptTemplate.from_messages([
        SystemMessage(content="Você é um especialista em criação de currículos."),
        HumanMessage(content=f"{structure_description}\n\nPLANO:\n{plan}\n\nCV ORIGINAL:\n{cv_text}\n\nVAGA:\n{job_desc}")
    ])

    chain = prompt | llm
    optimized_cv = chain.invoke({}).content

    # --- limpeza ---
    optimized_cv = re.sub(r"\*\*(.*?)\*\*", r"\1", optimized_cv)
    optimized_cv = re.sub(r"[*_`#>-]", "", optimized_cv)
    optimized_cv = re.sub(r"\n{3,}", "\n\n", optimized_cv.strip())

    # --- limitar para 500 palavras ---
    words = optimized_cv.split()
    if len(words) > 500:
        optimized_cv = " ".join(words[:500])

    # --- garantir 1 de cada seção ---
    sections = ["Profile", "Experiência", "Education", "Educação", "Habilidades", "Skills", "Projetos", "Certificações"]
    final_text = []
    seen_sections = set()
    for line in optimized_cv.split("\n"):
        for s in sections:
            if line.strip().startswith(s) and s in seen_sections:
                line = ""  # remove se repetido
            elif line.strip().startswith(s):
                seen_sections.add(s)
        if line.strip():
            final_text.append(line)
    optimized_cv = "\n".join(final_text)

    return {
        "optimized_cv": optimized_cv,
        "revision_needed": False,
        "review_feedback": feedback,
        "revision_count": revision_count
    }

def review_cv(state: CVAgentState):
    st.info("🔎 Avaliando qualidade do currículo...")
    optimized_cv = state["optimized_cv"]

    review_prompt = f"""
Você é um recrutador experiente.
Avalie o currículo adaptado quanto a clareza, relevância e adequação à vaga.
Se estiver pronto, responda apenas "OK".
Caso contrário, inicie com "REVISÃO:" e liste as melhorias necessárias.
---
{optimized_cv[:1000]}...
"""
    chain = ChatPromptTemplate.from_template(review_prompt) | llm
    review_result = chain.invoke({}).content.strip()
    result_keyword = review_result.split(":")[0].strip().upper()

    MAX_REVISIONS = 3
    if state.get("revision_count", 0) >= MAX_REVISIONS:
        st.warning("⚠️ Limite de revisões atingido.")
        return {"revision_needed": False, "review_feedback": "OK"}

    if "REVISÃO" in result_keyword:
        return {"revision_needed": True, "review_feedback": review_result}
    else:
        st.success("✅ Currículo final aprovado!")
        return {"revision_needed": False, "review_feedback": "OK"}

# ==========================================================
# IV. GRAFO
# ==========================================================

def should_continue(state: CVAgentState):
    if state["revision_needed"] and state.get("revision_count", 0) < 3:
        return "re_draft"
    else:
        return "publish"

def create_agent_graph():
    workflow = StateGraph(CVAgentState)
    workflow.add_node("evaluate", evaluate_compatibility)
    workflow.add_node("analyze", analyze_cv)
    workflow.add_node("optimize", optimize_cv)
    workflow.add_node("review", review_cv)

    workflow.set_entry_point("evaluate")
    workflow.add_edge("evaluate", "analyze")
    workflow.add_edge("analyze", "optimize")
    workflow.add_edge("optimize", "review")
    workflow.add_conditional_edges("review", should_continue, {"re_draft": "optimize", "publish": END})

    return workflow.compile()

# ==========================================================
# V. INTERFACE STREAMLIT
# ==========================================================

st.set_page_config(page_title="CV Optimizer Agent", layout="wide")
st.title("💼 Otimizador de Currículos Inteligente (.docx)")

if not LANGCHAIN_API_KEY or not GROQ_API_KEY:
    st.error("⚠️ As chaves `LANGCHAIN_API_KEY` e `GROQ_API_KEY` devem estar definidas no .env")
else:
    st.subheader("📋 Insira suas informações")

    uploaded_cv = st.file_uploader("Envie seu currículo (.docx)", type=["docx"])
    job_desc = st.text_area("🧾 Descrição da vaga", height=200)
    profile = st.text_area("👤 Perfil / Sumário Profissional", height=100)
    experiences = st.text_area("💼 Experiências", height=150)
    education = st.text_area("🎓 Educação / Formação", height=100)
    skills = st.text_area("🧠 Habilidades Técnicas", height=100)
    projects = st.text_area("🚀 Projetos", height=100)
    certifications = st.text_area("📜 Certificações", height=100)

    generate_btn = st.button("🚀 Avaliar e Otimizar Currículo")

    if generate_btn and uploaded_cv and job_desc.strip():
        doc = Document(uploaded_cv)
        cv_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""])

        with st.spinner("Executando agente de otimização..."):
            app = create_agent_graph()
            final_state = app.invoke({
                "cv_text": cv_text,
                "job_description": job_desc.strip(),
                "profile": profile,
                "experiences": experiences,
                "education": education,
                "skills": skills,
                "projects": projects,
                "certifications": certifications
            })

        compatibility = final_state.get("compatibility", 0)
        optimized_cv = final_state["optimized_cv"]

        st.markdown(f"### 🎯 Compatibilidade com a vaga: **{compatibility:.1f}%**")

        st.subheader("📝 Currículo Otimizado")
        st.text_area("Visualização do conteúdo gerado", value=optimized_cv, height=400)

        # ===== Criação do DOCX =====
        new_doc = Document()

        # --- Cabeçalho ---
        header_p = new_doc.add_paragraph()
        header_p.alignment = random.choice([0, 2])  # esquerda/direita

        linkedin_url = "https://www.linkedin.com/in/gabrielchaves"
        github_url = "https://github.com/gabrielchaves"
        site_url = "https://gabrielchaves.dev"
        email = "gabrielchavesmarques@hotmail.com"
        telefone = "+351 934663348"

        def add_hyperlink(paragraph, text, url):
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            part = paragraph.part
            r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), r_id)
            new_run = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            new_run.append(rPr)
            t = OxmlElement("w:t")
            t.text = text
            new_run.append(t)
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)

        # Nome do candidato em destaque
        name_p = new_doc.add_paragraph()
        name_run = name_p.add_run("Gabriel Chaves - Junior Fullstack Software Engineer")
        name_run.bold = True
        name_run.font.size = Pt(18)
        name_p.paragraph_format.space_after = Pt(8)

        # Links do portfólio
        links_p = new_doc.add_paragraph()
        add_hyperlink(links_p, "GitHub", github_url)
        links_p.add_run(" | ")
        add_hyperlink(links_p, "LinkedIn", linkedin_url)
        links_p.add_run(" | ")
        add_hyperlink(links_p, "Site", site_url)
        links_p.add_run(f" | Email: {email} | Telefone: {telefone}")
        links_p.paragraph_format.space_after = Pt(8)

        # Espaço para foto
        photo_run = new_doc.add_paragraph().add_run("[Espaço reservado para foto]")
        photo_run.italic = True
        photo_run.font.size = Pt(12)
        new_doc.add_paragraph().add_run("\n")

        # --- Conteúdo do CV ---
        bold_titles = [
            "Perfil", "Experiência", "Educação",
            "Habilidades Técnicas", "Linguagens de Programação",
            "Conceitos", "Projetos", "Certificações"
        ]

        for line in optimized_cv.split("\n"):
            if not line.strip():
                continue
            p = new_doc.add_paragraph()
            run = p.add_run(line.strip())
            if any(line.strip().startswith(title) for title in bold_titles):
                run.bold = True
                run.font.size = Pt(14)
            else:
                run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(4)

            # Quebra de página antes de Projetos e Certificações se estiverem muito longas
            if line.strip().startswith("Projetos") or line.strip().startswith("Certificações"):
                new_doc.add_page_break()

        # --- Links clicáveis nos projetos ---
        for paragraph in new_doc.paragraphs:
            match = re.search(r'(https?://github\.com/\S+)', paragraph.text)
            if match:
                url = match.group(1)
                paragraph.clear()
                add_hyperlink(paragraph, url, url)

        # --- Salvar DOCX para download ---
        bio = BytesIO()
        new_doc.save(bio)
        bio.seek(0)

        st.download_button(
            label="⬇️ Baixar Currículo Otimizado (.docx)",
            data=bio,
            file_name="CV_Otimizado.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# Ídeias
# O user dá o link do GH e Linkedin assim o programa n tem de descobrir no docx
# Implementar função ATS
# Melhorar formatação
# Não deixar aparecer CV duplicado